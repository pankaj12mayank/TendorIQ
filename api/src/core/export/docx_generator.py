"""Export Engine - DOCX Generator using python-docx"""

import io
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .schemas import ExportTemplate, ExportType
from .engine import get_template_manager

logger = logging.getLogger(__name__)


def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip('#')
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return RGBColor(r, g, b)


def set_cell_shading(cell, color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)


class DOCXGenerator:
    def __init__(self, template: Optional[ExportTemplate] = None):
        self.template = template or get_template_manager().get_default()
        self.primary = hex_to_rgb(self.template.primary_color)
        self.secondary = hex_to_rgb(self.template.secondary_color)
        self.accent = hex_to_rgb(self.template.accent_color)

    def _add_horizontal_line(self, doc: Document):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self.template.primary_color.lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_header_table(self, doc: Document, title: str, metadata: dict):
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_cell_shading(cell, self.template.primary_color)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(255, 255, 255)

        if metadata:
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for key, value in metadata.items():
                run = p2.add_run(f'{key}: {value}  ')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(220, 220, 220)

        doc.add_paragraph()

    def _add_section_header(self, doc: Document, title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = self.primary

    def _add_body_text(self, doc: Document, text: str):
        if not text:
            return
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5

            for line in para_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('- ') or line.startswith('* '):
                    p.add_run('\n• ' + line[2:])
                elif line.startswith(('1.', '2.', '3.', '4.', '5.')):
                    p.add_run('\n' + line)
                else:
                    p.add_run(line)
                p.add_run(' ')

            if p.text:
                p.text = p.text.rstrip()

    def _add_metadata_table(self, doc: Document, data: list):
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, (key, value) in enumerate(data):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(value)

            set_cell_shading(row.cells[0], '#f7fafc')

        doc.add_paragraph()

    def _add_checklist_table(self, doc: Document, items: list, include_due: bool = True):
        headers = ['Status', 'Item', 'Mandatory']
        if include_due:
            headers.append('Due Date')

        table = doc.add_table(rows=1 + len(items), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, self.template.primary_color.lstrip('#'))

        for i, item in enumerate(items):
            row = table.rows[i + 1]
            status = '✓' if item.get('is_submitted') else '○'
            row.cells[0].text = status
            row.cells[1].text = item.get('name', '')
            row.cells[2].text = 'Yes' if item.get('is_mandatory') else 'No'
            if include_due:
                due = item.get('due_date', '')
                row.cells[3].text = str(due)[:10] if due else 'N/A'

        doc.add_paragraph()

    def _add_risk_table(self, doc: Document, risks: list):
        table = doc.add_table(rows=1 + len(risks), cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        headers = ['Risk', 'Severity', 'Impact', 'Mitigation']
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, self.template.primary_color.lstrip('#'))

        severity_colors = {
            'critical': '#c53030',
            'high': '#dd6b20',
            'medium': '#d69e2e',
            'low': '#38a169',
        }

        for i, risk in enumerate(risks):
            row = table.rows[i + 1]
            row.cells[0].text = risk.get('title', '')
            severity = risk.get('severity', 'low').lower()
            row.cells[1].text = severity.upper()
            set_cell_shading(row.cells[1], severity_colors.get(severity, '#718096'))
            row.cells[2].text = risk.get('impact', 'medium').title()
            mitigation = risk.get('mitigation', '') or 'TBD'
            if len(mitigation) > 60:
                mitigation = mitigation[:57] + '...'
            row.cells[3].text = mitigation

        doc.add_paragraph()

    def generate_proposal_docx(self, data: dict) -> bytes:
        doc = Document()
        doc.sections[0].page_width = Cm(21)
        doc.sections[0].page_height = Cm(29.7)
        doc.sections[0].left_margin = Cm(2)
        doc.sections[0].right_margin = Cm(2)
        doc.sections[0].top_margin = Cm(2.5)
        doc.sections[0].bottom_margin = Cm(2.5)

        metadata = {}
        if data.get('status'):
            metadata['Status'] = data['status']
        if data.get('version'):
            metadata['Version'] = data['version']
        if data.get('created_at'):
            metadata['Generated'] = data['created_at']
        if data.get('organization'):
            metadata['Organization'] = data['organization']

        self._add_header_table(doc, data.get('title', 'Proposal Document'), metadata)
        self._add_horizontal_line(doc)

        if data.get('summary'):
            self._add_section_header(doc, 'Executive Summary')
            self._add_body_text(doc, data['summary'])

        for section in data.get('sections', []):
            self._add_section_header(doc, section.get('title', ''))
            self._add_body_text(doc, section.get('content', ''))
            if section.get('word_count'):
                p = doc.add_paragraph()
                run = p.add_run(f'Word count: {section["word_count"]}')
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = self.secondary

        doc.add_page_break()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('--- END OF DOCUMENT ---')
        run.font.size = Pt(10)
        run.font.color.rgb = self.secondary

        if self.template.show_timestamp:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'Generated on {datetime.utcnow().strftime("%d %b %Y, %H:%M UTC")}')
            run.font.size = Pt(8)
            run.font.color.rgb = self.secondary

        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()
        buffer.close()
        return content

    def generate_checklist_docx(self, data: dict) -> bytes:
        doc = Document()
        doc.sections[0].page_width = Cm(21)
        doc.sections[0].page_height = Cm(29.7)
        doc.sections[0].left_margin = Cm(2)
        doc.sections[0].right_margin = Cm(2)
        doc.sections[0].top_margin = Cm(2.5)
        doc.sections[0].bottom_margin = Cm(2.5)

        metadata = [
            ['Progress', f"{data.get('completion_percentage', 0):.1f}%"],
            ['Overall Score', f"{data.get('score', {}).get('overall_score', 0):.1f}/100"],
            ['Mandatory', f"{data.get('mandatory_completed', 0)}/{data.get('mandatory_total', 0)}"],
        ]
        self._add_header_table(doc, data.get('name', 'Compliance Checklist'), dict(metadata))
        self._add_horizontal_line(doc)

        if data.get('description'):
            self._add_body_text(doc, data['description'])
            doc.add_paragraph()

        for section in data.get('sections', []):
            self._add_section_header(doc, section.get('name', ''))

            section_items = section.get('items', [])
            self._add_checklist_table(doc, section_items)

        if data.get('submission_steps'):
            self._add_section_header(doc, 'Submission Steps')
            for i, step in enumerate(data.get('submission_steps', []), 1):
                p = doc.add_paragraph()
                status = '✓' if step.get('is_completed') else '○'
                run = p.add_run(f'{i}. {status} {step.get("name", "")}')
                run.bold = True
                if step.get('estimated_duration_minutes'):
                    run2 = p.add_run(f" ({step['estimated_duration_minutes']} min)")
                    run2.font.size = Pt(9)
                    run2.font.italic = True

        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Generated on {datetime.utcnow().strftime("%d %b %Y, %H:%M UTC")}')
        run.font.size = Pt(8)
        run.font.color.rgb = self.secondary

        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()
        buffer.close()
        return content

    def generate_risk_analysis_docx(self, data: dict) -> bytes:
        doc = Document()
        doc.sections[0].page_width = Cm(21)
        doc.sections[0].page_height = Cm(29.7)
        doc.sections[0].left_margin = Cm(2)
        doc.sections[0].right_margin = Cm(2)
        doc.sections[0].top_margin = Cm(2.5)
        doc.sections[0].bottom_margin = Cm(2.5)

        metadata = {}
        if data.get('tender_id'):
            metadata['Tender ID'] = data['tender_id']
        if data.get('overall_risk_score') is not None:
            score = data['overall_risk_score']
            severity = self._get_risk_severity(score)
            metadata['Risk Score'] = f'{score:.1f}/100 ({severity})'
        if data.get('analysis_date'):
            metadata['Analysis Date'] = data['analysis_date']

        self._add_header_table(doc, data.get('title', 'Risk Analysis Report'), metadata)
        self._add_horizontal_line(doc)

        if data.get('summary'):
            self._add_section_header(doc, 'Executive Summary')
            self._add_body_text(doc, data['summary'])

        for category in data.get('risk_categories', []):
            self._add_section_header(doc, category.get('name', 'Category'))

            risks = category.get('risks', [])
            if risks:
                self._add_risk_table(doc, risks)

        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Generated on {datetime.utcnow().strftime("%d %b %Y, %H:%M UTC")}')
        run.font.size = Pt(8)
        run.font.color.rgb = self.secondary

        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()
        buffer.close()
        return content

    def _get_risk_severity(self, score: float) -> str:
        if score >= 90:
            return 'Critical'
        elif score >= 70:
            return 'High'
        elif score >= 40:
            return 'Medium'
        else:
            return 'Low'

    def generate_generic_docx(self, data: dict, doc_type: ExportType = ExportType.CUSTOM) -> bytes:
        if doc_type == ExportType.PROPOSAL:
            return self.generate_proposal_docx(data)
        elif doc_type == ExportType.CHECKLIST:
            return self.generate_checklist_docx(data)
        elif doc_type == ExportType.RISK_ANALYSIS:
            return self.generate_risk_analysis_docx(data)
        else:
            return self.generate_custom_docx(data)

    def generate_custom_docx(self, data: dict) -> bytes:
        doc = Document()
        doc.sections[0].page_width = Cm(21)
        doc.sections[0].page_height = Cm(29.7)
        doc.sections[0].left_margin = Cm(2)
        doc.sections[0].right_margin = Cm(2)
        doc.sections[0].top_margin = Cm(2.5)
        doc.sections[0].bottom_margin = Cm(2.5)

        title = data.get('title') or data.get('name') or 'Document'
        self._add_header_table(doc, title, {})

        meta_data = []
        for key, value in data.items():
            if key in ('title', 'name', 'sections', 'items', 'risks', 'content', 'summary'):
                continue
            meta_data.append((key.replace('_', ' ').title(), str(value)))

        if meta_data:
            self._add_metadata_table(doc, meta_data)

        self._add_horizontal_line(doc)

        if 'sections' in data:
            for section in data['sections']:
                self._add_section_header(doc, section.get('title', ''))
                content = section.get('content', '') or section.get('description', '')
                self._add_body_text(doc, content)

        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'Generated on {datetime.utcnow().strftime("%d %b %Y, %H:%M UTC")}')
        run.font.size = Pt(8)
        run.font.color.rgb = self.secondary

        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()
        buffer.close()
        return content


docx_generator = DOCXGenerator()


def get_docx_generator(template: Optional[ExportTemplate] = None) -> DOCXGenerator:
    return DOCXGenerator(template)