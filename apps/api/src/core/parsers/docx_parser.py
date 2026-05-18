"""DOCX Parser using python-docx"""

import io
import logging
import re
from datetime import datetime
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt

from .base import (
    BaseParser, ParsedDocument, DocumentMetadata,
    PageSegment, Section, ParsingError
)
from ..logging import get_logger

logger = get_logger('docx_parser')


class DOCXParser(BaseParser):
    supported_extensions = ['docx', 'docm']

    def __init__(self):
        pass

    async def parse(self, file_bytes: bytes, file_name: str, document_id: str) -> ParsedDocument:
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
        except Exception as e:
            raise ParsingError(f'Failed to open DOCX: {str(e)}')

        metadata = await self._extract_metadata(doc, file_bytes, file_name)
        full_text, pages = await self._extract_pages(doc)
        full_text = self.normalize_text(full_text)

        sections = self._extract_sections(doc)
        links = self._extract_links(doc)
        tables = self._extract_tables(doc)

        return ParsedDocument(
            document_id=document_id,
            file_name=file_name,
            file_type='docx',
            metadata=metadata,
            full_text=full_text,
            pages=pages,
            sections=sections,
            tables=tables,
            images=[],
            links=links,
            confidence_score=self._calculate_confidence(full_text, pages),
        )

    async def extract_metadata(self, file_bytes: bytes) -> DocumentMetadata:
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            return await self._extract_metadata(doc, file_bytes, '')
        except Exception:
            return DocumentMetadata()

    async def extract_text(self, file_bytes: bytes) -> str:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        return self.normalize_text('\n'.join(paragraphs))

    async def _extract_metadata(self, doc: DocxDocument, file_bytes: bytes, file_name: str) -> DocumentMetadata:
        props = doc.core_properties

        keywords = []
        if props.keywords:
            keywords = [k.strip() for k in str(props.keywords).split(',') if k.strip()]

        text_parts = [p.text for p in doc.paragraphs]
        word_count = len(' '.join(text_parts).split())

        return DocumentMetadata(
            title=props.title or file_name.replace('.docx', '').replace('.docm', ''),
            author=str(props.author) if props.author else None,
            creator=str(props.author) if props.author else None,
            subject=props.subject or None,
            description=props.subject or None,
            keywords=keywords,
            creation_date=props.created,
            modification_date=props.modified,
            page_count=len(doc.sections),
            word_count=word_count,
            language='en',
            file_size=len(file_bytes),
            is_encrypted=False,
            is_form=len(doc.form_fields) > 0,
        )

    async def _extract_pages(self, doc: DocxDocument) -> tuple[str, list[PageSegment]]:
        pages = []
        all_text = []
        char_pos = 0

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                char_pos += 1
                continue

            is_heading = para.style.name.startswith('Heading') or para.style.name.startswith('Title')
            heading_level = None

            if 'Title' in para.style.name:
                heading_level = 0
            elif para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                    heading_level = level
                except:
                    heading_level = 1
            else:
                font_size = None
                for run in para.runs:
                    if run.font.size:
                        font_size = run.font.size.pt
                        break

                if font_size and font_size >= 14:
                    is_heading = True
                    heading_level = 1
                elif font_size and font_size >= 12:
                    is_heading = True
                    heading_level = 2

            is_bold = any(run.bold for run in para.runs if run.text.strip())
            if is_bold and len(text.strip()) < 100 and not text.endswith(':'):
                is_heading = True
                if heading_level is None:
                    heading_level = 2

            seg = PageSegment(
                page_number=1,
                text=text.strip(),
                start_char=char_pos,
                end_char=char_pos + len(text.strip()),
                is_heading=is_heading,
                heading_level=heading_level,
                section_title=text.strip() if is_heading else None,
                confidence=1.0,
                font_info=None,
            )

            pages.append(seg)
            all_text.append(text.strip())
            char_pos += len(text.strip()) + 1

        return '\n'.join(all_text), pages

    def _extract_sections(self, doc: DocxDocument) -> list[Section]:
        sections = []
        current_section = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = para.style.name.startswith('Heading') or para.style.name.startswith('Title')
            heading_level = 1

            if 'Title' in para.style.name:
                heading_level = 0
            elif para.style.name.startswith('Heading'):
                try:
                    heading_level = int(para.style.name.replace('Heading ', ''))
                except:
                    heading_level = 1

            if is_heading or (len(text) < 100 and any(r.bold for r in para.runs if r.text.strip())):
                if current_section:
                    current_section.content = '\n'.join(current_content).strip()
                    current_section.word_count = len(current_section.content.split())
                    sections.append(current_section)

                current_section = Section(
                    title=text,
                    level=heading_level,
                    start_page=1,
                    end_page=1,
                    content='',
                )
                current_content = []
            else:
                if current_section:
                    current_content.append(text)

        if current_section:
            current_section.content = '\n'.join(current_content).strip()
            current_section.word_count = len(current_section.content.split())
            sections.append(current_section)

        return sections

    def _extract_links(self, doc: DocxDocument) -> list[dict]:
        links = []
        for rel in doc.part.rels.values():
            if 'hyperlink' in rel.reltype:
                links.append({
                    'uri': rel.target_ref,
                    'type': 'hyperlink',
                })
        return links

    def _extract_tables(self, doc: DocxDocument) -> list[dict]:
        tables = []
        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0

            cells_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                cells_text.append(row_text)

            tables.append({
                'rows': rows,
                'columns': cols,
                'cells': cells_text,
            })

        return tables

    def _calculate_confidence(self, full_text: str, pages: list[PageSegment]) -> float:
        if not full_text:
            return 0.0

        text_ratio = min(1.0, len(full_text) / 100)

        heading_count = sum(1 for p in pages if p.is_heading)
        heading_ratio = min(1.0, heading_count / max(1, len(pages) * 0.2))

        avg_word_len = sum(len(w) for w in full_text.split()) / max(1, len(full_text.split()))
        quality_ratio = min(1.0, avg_word_len / 4)

        confidence = (text_ratio * 0.3 + heading_ratio * 0.4 + quality_ratio * 0.3)
        return round(confidence, 3)


docx_parser = DOCXParser()